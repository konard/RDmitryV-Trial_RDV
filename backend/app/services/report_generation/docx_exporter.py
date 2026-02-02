"""DOCX report exporter."""

from typing import Dict, Any, List, Optional
from pathlib import Path
import io

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .gost_formatter import GOSTFormatter


class DOCXExporter:
    """
    DOCX report exporter.

    Exports report to DOCX format with GOST 7.32-2017 formatting.
    """

    def __init__(self):
        self.formatter = GOSTFormatter()
        self.document = None
        self.section_numbers = {}
        self.figure_counter = 0
        self.table_counter = 0
        self.reference_counter = 0

    def create_document(self, report_data: Dict[str, Any]) -> bytes:
        """
        Create DOCX document from report data.

        Args:
            report_data: Dictionary with report content

        Returns:
            DOCX file bytes
        """
        self.document = Document()
        self._setup_document_styles()
        self._setup_page_settings()

        # Generate document sections
        self._add_title_page(report_data.get("title_page", {}))
        self._add_abstract(report_data.get("abstract", {}))
        self._add_table_of_contents(report_data.get("sections", []))
        self._add_introduction(report_data.get("introduction", {}))
        self._add_main_sections(report_data.get("main_sections", []))
        self._add_conclusion(report_data.get("conclusion", {}))
        self._add_bibliography(report_data.get("bibliography", []))
        self._add_appendices(report_data.get("appendices", []))

        # Save to bytes
        return self._save_to_bytes()

    def _setup_document_styles(self) -> None:
        """Setup document styles according to GOST."""
        styles = self.document.styles

        # Normal style
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = self.formatter.page_settings.font_name
        normal_font.size = Pt(self.formatter.page_settings.font_size)

        paragraph_format = normal_style.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph_format.first_line_indent = Cm(1.25)

        # Heading styles
        self._create_heading_style('GOST Heading 0', 14, True, WD_ALIGN_PARAGRAPH.CENTER, True)
        self._create_heading_style('GOST Heading 1', 14, True, WD_ALIGN_PARAGRAPH.LEFT, False)
        self._create_heading_style('GOST Heading 2', 14, True, WD_ALIGN_PARAGRAPH.LEFT, False)
        self._create_heading_style('GOST Heading 3', 14, False, WD_ALIGN_PARAGRAPH.LEFT, False)

    def _create_heading_style(
        self,
        name: str,
        size: int,
        bold: bool,
        alignment: WD_ALIGN_PARAGRAPH,
        uppercase: bool
    ) -> None:
        """Create custom heading style."""
        styles = self.document.styles

        try:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            # Style already exists
            style = styles[name]

        font = style.font
        font.name = self.formatter.page_settings.font_name
        font.size = Pt(size)
        font.bold = bold

        paragraph_format = style.paragraph_format
        paragraph_format.alignment = alignment
        paragraph_format.space_before = Pt(12)
        paragraph_format.space_after = Pt(12)
        paragraph_format.keep_with_next = True

    def _setup_page_settings(self) -> None:
        """Setup page settings according to GOST."""
        section = self.document.sections[0]

        # Page size (A4)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

        # Margins
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

        # Page numbering
        self._add_page_numbers(section)

    def _add_page_numbers(self, section) -> None:
        """Add page numbers to footer."""
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add page number field
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)

    def _add_title_page(self, data: Dict[str, Any]) -> None:
        """Add title page."""
        # Organization name
        p = self.document.add_paragraph(data.get("organization", ""))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(14)
        p.runs[0].font.bold = True

        # Add spacing
        for _ in range(10):
            self.document.add_paragraph()

        # Report type
        p = self.document.add_paragraph(data.get("report_type", "").upper())
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(16)
        p.runs[0].font.bold = True

        self.document.add_paragraph()

        # Title
        p = self.document.add_paragraph(data.get("title", ""))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(18)
        p.runs[0].font.bold = True

        # Add spacing
        for _ in range(10):
            self.document.add_paragraph()

        # Date and location
        p = self.document.add_paragraph(f"{data.get('date', '')}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(14)

        # Page break
        self.document.add_page_break()

    def _add_abstract(self, data: Dict[str, Any]) -> None:
        """Add abstract section."""
        # Heading
        p = self.document.add_paragraph("РЕФЕРАТ", style='GOST Heading 0')

        # Statistics
        stats = data.get("statistics", {})
        stats_text = (
            f"Отчет содержит {stats.get('pages', 0)} страниц, "
            f"{stats.get('figures', 0)} рисунков, "
            f"{stats.get('tables', 0)} таблиц, "
            f"{stats.get('sources', 0)} источников, "
            f"{stats.get('appendices', 0)} приложений."
        )
        self.document.add_paragraph(stats_text)

        # Keywords
        keywords = data.get("keywords", [])
        if keywords:
            keywords_text = "КЛЮЧЕВЫЕ СЛОВА: " + ", ".join(keywords).upper()
            p = self.document.add_paragraph(keywords_text)
            p.runs[0].font.bold = True

        self.document.add_paragraph()

        # Summary
        summary = data.get("summary", "")
        self.document.add_paragraph(summary)

        # Page break
        self.document.add_page_break()

    def _add_table_of_contents(self, sections: List[Dict[str, Any]]) -> None:
        """Add table of contents."""
        # Heading
        p = self.document.add_paragraph("СОДЕРЖАНИЕ", style='GOST Heading 0')

        # TOC entries
        toc_items = [
            ("ВВЕДЕНИЕ", 3),
        ]

        # Add main sections
        for i, section in enumerate(sections, 1):
            toc_items.append((f"{i} {section.get('title', '').upper()}", section.get('page', i + 3)))

            # Add subsections if any
            for j, subsection in enumerate(section.get('subsections', []), 1):
                toc_items.append((f"  {i}.{j} {subsection.get('title', '')}", subsection.get('page', 0)))

        toc_items.extend([
            ("ЗАКЛЮЧЕНИЕ", len(sections) + 5),
            ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", len(sections) + 6),
        ])

        # Add TOC entries
        for title, page in toc_items:
            p = self.document.add_paragraph()
            p.add_run(title)
            p.add_run("." * (100 - len(title) - len(str(page))))
            p.add_run(str(page))

        # Page break
        self.document.add_page_break()

    def _add_introduction(self, data: Dict[str, Any]) -> None:
        """Add introduction section."""
        # Heading
        p = self.document.add_paragraph("ВВЕДЕНИЕ", style='GOST Heading 0')

        # Content
        text = data.get("text", "")
        self.document.add_paragraph(text)

        # Page break
        self.document.add_page_break()

    def _add_main_sections(self, sections: List[Dict[str, Any]]) -> None:
        """Add main sections."""
        for i, section in enumerate(sections, 1):
            self.section_numbers[section.get("id", "")] = i

            # Section heading
            title = f"{i} {section.get('title', '').upper()}"
            p = self.document.add_paragraph(title, style='GOST Heading 1')

            # Section content
            content = section.get("content", "")
            if content:
                self.document.add_paragraph(content)

            # Add subsections
            for j, subsection in enumerate(section.get("subsections", []), 1):
                # Subsection heading
                subtitle = f"{i}.{j} {subsection.get('title', '')}"
                p = self.document.add_paragraph(subtitle, style='GOST Heading 2')

                # Subsection content
                subcontent = subsection.get("content", "")
                if subcontent:
                    self.document.add_paragraph(subcontent)

                # Add figures
                for figure in subsection.get("figures", []):
                    self._add_figure(figure)

                # Add tables
                for table in subsection.get("tables", []):
                    self._add_table(table)

            # Page break after each section
            self.document.add_page_break()

    def _add_figure(self, figure_data: Dict[str, Any]) -> None:
        """Add figure to document."""
        self.figure_counter += 1

        # Add image if available
        image_data = figure_data.get("image_bytes")
        if image_data:
            image_stream = io.BytesIO(image_data)
            self.document.add_picture(image_stream, width=Inches(6))

        # Add caption
        caption = self.formatter.format_figure_caption(
            self.figure_counter,
            figure_data.get("title", "")
        )
        p = self.document.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.document.add_paragraph()

    def _add_table(self, table_data: Dict[str, Any]) -> None:
        """Add table to document."""
        self.table_counter += 1

        # Add caption
        caption = self.formatter.format_table_caption(
            self.table_counter,
            table_data.get("title", "")
        )
        p = self.document.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Create table
        data = table_data.get("data", [])
        headers = table_data.get("headers", [])

        if data and headers:
            table = self.document.add_table(rows=len(data) + 1, cols=len(headers))
            table.style = 'Table Grid'

            # Add headers
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = str(header)
                cell.paragraphs[0].runs[0].font.bold = True

            # Add data
            for row_idx, row_data in enumerate(data, 1):
                for col_idx, value in enumerate(row_data):
                    table.rows[row_idx].cells[col_idx].text = str(value)

        self.document.add_paragraph()

    def _add_conclusion(self, data: Dict[str, Any]) -> None:
        """Add conclusion section."""
        # Heading
        p = self.document.add_paragraph("ЗАКЛЮЧЕНИЕ", style='GOST Heading 0')

        # Content
        text = data.get("text", "")
        self.document.add_paragraph(text)

        # Page break
        self.document.add_page_break()

    def _add_bibliography(self, bibliography: List[Dict[str, Any]]) -> None:
        """Add bibliography section."""
        # Heading
        p = self.document.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", style='GOST Heading 0')

        # Bibliography entries
        for i, entry in enumerate(bibliography, 1):
            # Format entry
            entry_text = self._format_bibliography_entry(entry)

            # Add entry with number
            p = self.document.add_paragraph(f"{i}. {entry_text}")
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.hanging_indent = Cm(-1.0)

        # Page break
        self.document.add_page_break()

    def _format_bibliography_entry(self, entry: Dict[str, Any]) -> str:
        """Format bibliography entry."""
        entry_type = entry.get("type", "website")

        if entry_type == "website":
            title = entry.get("title", "")
            url = entry.get("url", "")
            access_date = entry.get("access_date", "")

            text = f"{title}"
            if url:
                text += f". – URL: {url}"
            if access_date:
                text += f" (дата обращения: {access_date})"

            return text

        # Default format
        return entry.get("title", "")

    def _add_appendices(self, appendices: List[Dict[str, Any]]) -> None:
        """Add appendices."""
        if not appendices:
            return

        for appendix in appendices:
            # Heading
            p = self.document.add_paragraph(f"ПРИЛОЖЕНИЕ {appendix.get('letter', 'А')}", style='GOST Heading 0')
            p = self.document.add_paragraph(appendix.get('title', ''))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Content
            content = appendix.get("content", "")
            self.document.add_paragraph(content)

            # Page break
            self.document.add_page_break()

    def _save_to_bytes(self) -> bytes:
        """Save document to bytes."""
        buffer = io.BytesIO()
        self.document.save(buffer)
        buffer.seek(0)
        return buffer.read()
